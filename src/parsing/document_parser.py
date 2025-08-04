"""
Document parsing utilities for PDF and DOCX files
"""

import re
import io
from typing import Dict, List, Optional, Set, Match
import PyPDF2
import pdfplumber
from docx import Document
from pydantic import BaseModel, Field


class ParsedDocument(BaseModel):
    """Type-safe parsed document content"""
    text: str = Field(..., description="Extracted text content")
    filename: str = Field(..., description="Original filename")
    file_type: str = Field(..., description="File type (pdf, docx)")
    page_count: Optional[int] = Field(default=None, description="Number of pages")


class ExtractedAnswers(BaseModel):
    """Type-safe extracted answers from document"""
    answers: Dict[str, str] = Field(..., description="Question ID to answer mapping")
    question_count: int = Field(..., gt=0, description="Expected number of questions")
    student_name: str = Field(..., description="Extracted student name")


class DocumentParser:
    """Handles parsing of PDF and DOCX documents"""
    
    def __init__(self) -> None:
        self.answer_patterns: List[str] = [
            r'^(\d+)[\.\)]\s*',  # 1. or 1)
            r'^Q(\d+)[\.\)]\s*',  # Q1. or Q1)
            r'^Question\s*(\d+)[\.\)]\s*',  # Question 1. or Question 1)
            r'^Ans[\.\s]*(\d+)[\.\)]\s*',  # Ans. 1 or Ans 1)
        ]
    
    def parse_pdf(self, file_content: bytes) -> str:
        """Parse PDF file and extract text"""
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_content: List[str] = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                
                if text_content:
                    return '\n'.join(text_content)
            
            # Fallback to PyPDF2
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_content = []
            
            for page in pdf_reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text_content.append(extracted_text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            raise
    
    def parse_docx(self, file_content: bytes) -> str:
        """Parse DOCX file and extract text"""
        try:
            doc = Document(io.BytesIO(file_content))
            text_content: List[str] = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            raise
    
    def parse_document(self, file_content: bytes, filename: str) -> str:
        """Parse document based on file extension"""
        file_extension: str = filename.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self.parse_pdf(file_content)
        elif file_extension in ['docx', 'doc']:
            return self.parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def extract_answers_from_text(self, text: str, question_count: int) -> Dict[str, str]:
        """Extract individual answers from parsed text"""
        answers: Dict[str, str] = {}
        lines: List[str] = text.split('\n')
        current_answer: Optional[str] = None
        current_content: List[str] = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line starts with answer pattern
            answer_match: Optional[Match[str]] = None
            for pattern in self.answer_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    answer_match = match
                    break
            
            if answer_match:
                # Save previous answer if exists
                if current_answer and current_content:
                    answers[current_answer] = '\n'.join(current_content).strip()
                
                # Start new answer
                answer_num: str = answer_match.group(1)
                current_answer = f"Q{answer_num}"
                
                # Get content after the answer number
                remaining_text: str = line[answer_match.end():].strip()
                current_content = [remaining_text] if remaining_text else []
            else:
                # Continue current answer
                if current_answer:
                    current_content.append(line)
        
        # Save last answer
        if current_answer and current_content:
            answers[current_answer] = '\n'.join(current_content).strip()
        
        # Fill missing answers and guess numbering
        complete_answers: Dict[str, str] = self._fill_missing_answers(answers, question_count)
        
        return complete_answers
    
    def _fill_missing_answers(self, answers: Dict[str, str], expected_count: int) -> Dict[str, str]:
        """Fill missing answer numbers by guessing from existing pattern"""
        # Extract existing question numbers
        existing_nums: Set[int] = set()
        for key in answers.keys():
            if key.startswith('Q'):
                try:
                    num: int = int(key[1:])
                    existing_nums.add(num)
                except ValueError:
                    continue
        
        # Create complete answer set
        complete_answers: Dict[str, str] = {}
        
        for i in range(1, expected_count + 1):
            question_key: str = f"Q{i}"
            
            if i in existing_nums:
                # Use existing answer
                complete_answers[question_key] = answers.get(question_key, "")
            else:
                # Assign unmatched content or empty
                unassigned_content: Optional[str] = None
                for _, content in answers.items():
                    if content not in complete_answers.values():
                        unassigned_content = content
                        break
                
                complete_answers[question_key] = unassigned_content or ""
        
        return complete_answers
    
    def extract_student_name_from_filename(self, filename: str) -> str:
        """Extract student name from filename"""
        # Remove file extension
        name: str = filename.rsplit('.', 1)[0]
        
        # Clean up common patterns
        name = re.sub(r'[_\-]+', ' ', name)  # Replace underscores/hyphens with spaces
        name = re.sub(r'\s+', ' ', name)  # Normalize whitespace
        name = name.strip()
        
        # Capitalize words
        name = ' '.join(word.capitalize() for word in name.split())
        
        return name


# Global parser instance
document_parser = DocumentParser()
